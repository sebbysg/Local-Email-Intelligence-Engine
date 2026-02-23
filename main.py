import win32com.client
import ollama
import datetime
import re
import os
import requests
import logging
from docx import Document
from docx.shared import Pt
from dotenv import load_dotenv

# --- INITIALIZATION ---
load_dotenv()

# Configuration from .env
TELEGRAM_TOKEN = os.getenv("TELEGRAM_TOKEN")
TELEGRAM_CHAT_ID = os.getenv("TELEGRAM_CHAT_ID")
OUTPUT_DIR = os.getenv("OUTPUT_DIR")
MODEL_NAME = "gemma3:4b"
LOOKBACK_DAYS = 7

# Logging Setup
logging.basicConfig(
    filename='pipeline.log',
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)

def clean_body(text):
    """Strip signatures and email thread history."""
    patterns = [r"From:", r"-----Original Message-----", r"________________________________", r"On.*wrote:"]
    for pattern in patterns:
        text = re.split(pattern, text, flags=re.IGNORECASE)[0]
    return text.strip()

def get_pmo_data(days_back=7):
    """Fetches Inbox and Sent items, identifying 'Ghosted' emails."""
    logging.info("Accessing Outlook MAPI...")
    outlook = win32com.client.Dispatch("Outlook.Application").GetNamespace("MAPI")
    
    filter_date = (datetime.date.today() - datetime.timedelta(days=days_back)).strftime("%m/%d/%Y %H:%M %p")
    
    inbox = outlook.GetDefaultFolder(6).Items.Restrict(f"[ReceivedTime] >= '{filter_date}'")
    sent = outlook.GetDefaultFolder(5).Items.Restrict(f"[SentOn] >= '{filter_date}'")
    
    # Track replies via ConversationID
    replied_ids = {msg.ConversationID for msg in inbox}
    
    email_data = []
    
    # Process Sent Items
    for i, msg in enumerate(sent):
        try:
            email_data.append({
                "id": f"S{i+1}",
                "type": "SENT",
                "entry_id": msg.EntryID,
                "target": msg.To,
                "subject": msg.Subject,
                "body": clean_body(msg.Body),
                "is_waiting": msg.ConversationID not in replied_ids
            })
        except: continue

    # Process Inbox
    for i, msg in enumerate(inbox):
        try:
            email_data.append({
                "id": f"I{i+1}",
                "type": "INBOX",
                "entry_id": msg.EntryID,
                "target": msg.SenderName,
                "subject": msg.Subject,
                "body": clean_body(msg.Body),
                "is_waiting": False
            })
        except: continue
        
    return email_data

def generate_summary(data):
    """Sends context to Gemma 3 with PMO instructions."""
    logging.info(f"Querying {MODEL_NAME}...")
    
    context_list = []
    for e in data:
        status = "[WAITING FOR REPLY]" if e['is_waiting'] else ""
        context_list.append(f"SOURCE [{e['id']}] {e['type']} {status}\nFrom/To: {e['target']}\nSubj: {e['subject']}\nBody: {e['body'][:800]}")
    
    context_str = "\n---\n".join(context_list)

    prompt = f"""
    You are an IT Project Manager's automated reporting engine. Analyze the following emails.
    
    STRUCTURE:
    1. **Critical Blockers & Sentiment**: Note frustrations or stalls.
    2. **Unanswered Follow-ups**: List SENT emails with [WAITING FOR REPLY] status.
    3. **Weekly Progress**: High-level achievements.

    RULES:
    - NO TABLES. Use bullet points.
    - Quote sources like [SOURCE X].
    - NO conversational fluff or follow-up questions.
    
    DATA:
    {context_str}
    """

    response = ollama.chat(model=MODEL_NAME, messages=[
        {'role': 'system', 'content': 'You are a silent, formal PM reporting tool.'},
        {'role': 'user', 'content': prompt}
    ])
    return response['message']['content']

def save_report(summary, data):
    """Saves to Word with a detailed EntryID Audit Table."""
    if not os.path.exists(OUTPUT_DIR):
        os.makedirs(OUTPUT_DIR)
        
    doc = Document()
    doc.add_heading(f'PM Intelligence Report: {datetime.date.today()}', 0)
    doc.add_paragraph(summary)
    
    doc.add_page_break()
    doc.add_heading('Audit Trail: EntryID Reference', level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Ref'
    hdr_cells[1].text = 'Subject'
    hdr_cells[2].text = 'MAPI EntryID'

    for e in data:
        row = table.add_row().cells
        row[0].text = e['id']
        row[1].text = e['subject'][:50]
        # EntryIDs are long; use small font
        run = row[2].paragraphs[0].add_run(e['entry_id'])
        run.font.size = Pt(7)

    path = os.path.join(OUTPUT_DIR, f"Summary_{datetime.date.today()}.docx")
    doc.save(path)
    logging.info(f"Report saved to {path}")
    return path

def send_telegram(text):
    """Sends summary to Telegram via Bot API."""
    url = f"https://api.telegram.org/bot{TELEGRAM_TOKEN}/sendMessage"
    payload = {"chat_id": TELEGRAM_CHAT_ID, "text": f"📅 *Daily Update*\n\n{text[:4000]}", "parse_mode": "Markdown"}
    try:
        r = requests.post(url, data=payload)
        r.raise_for_status()
        logging.info("Telegram notification successful.")
    except Exception as e:
        logging.error(f"Telegram Failed: {e}")

def main():
    try:
        if datetime.datetime.today().weekday() >= 5:
            logging.info("Weekend detected. Skipping execution.")
            return

        email_data = get_pmo_data(LOOKBACK_DAYS)
        if not email_data:
            logging.info("No emails found in lookback period.")
            return

        summary = generate_summary(email_data)
        save_report(summary, email_data)
        send_telegram(summary)
        print("Workflow complete. Check pipeline.log for details.")
        
    except Exception as e:
        logging.critical(f"Pipeline crashed: {e}")

if __name__ == "__main__":
    main()